import os
import io
import re
import logging
import tempfile
import asyncio
import sqlite3
import threading
import requests
import zipfile
from datetime import datetime, timedelta
from io import BytesIO

import fitz
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageEnhance
import pytesseract
from PyPDF2 import PdfReader, PdfWriter
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.error import Forbidden, BadRequest
from telegram.ext import (
    Application, CommandHandler, MessageHandler, CallbackQueryHandler,
    filters, ContextTypes,
)
from flask import Flask, send_from_directory, request, jsonify, send_file
from flask_cors import CORS

# =============== LOGGING ===============
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# =============== KONFIGURATSIYA ===============
TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
if not TOKEN:
    raise RuntimeError("TELEGRAM_BOT_TOKEN environment variable topilmadi.")

ADMIN_ID = int(os.getenv("ADMIN_ID", "123456789"))
ADMIN_USERNAME = os.getenv("ADMIN_USERNAME", "@yoldoshev_3")
BOT_NAME = "📄 Professional File Converter"
MAX_FILE_SIZE = 50 * 1024 * 1024
OCR_LANGS = os.getenv("OCR_LANGS", "uzb+rus+eng")
DB_PATH = os.getenv("DB_PATH", "bot_stats.db")
PDF_TIMEOUT_SECONDS = int(os.getenv("PDF_TIMEOUT_SECONDS", "240"))
PRO_FREE_UNTIL = datetime(2026, 7, 14)

# =============== FLASK APP ===============
flask_app = Flask(__name__, static_folder='webapp', static_url_path='')
CORS(flask_app)

@flask_app.route('/webapp')
def serve_webapp():
    return send_from_directory('webapp', 'index.html')

@flask_app.route('/pro')
def serve_pro():
    return send_from_directory('webapp', 'pro.html')

@flask_app.route('/')
def home():
    return send_from_directory('webapp', 'index.html')

@flask_app.route('/api/convert', methods=['POST'])
def api_convert():
    try:
        file = request.files.get('file')
        if not file:
            return jsonify({'error': 'Fayl topilmadi'}), 400
        content = file.read()
        ext = file.filename.split('.')[-1].lower()
        output_name = file.filename.rsplit('.', 1)[0] + '.docx'
        if ext == 'pdf':
            output = pdf_to_word(content, file.filename)
        elif ext in ['xlsx', 'xls', 'csv']:
            output = excel_to_word(content, ext, file.filename)
        elif ext in ['html', 'htm']:
            output = html_to_word(content, file.filename)
        elif ext == 'epub':
            output = epub_to_word(content, file.filename)
        else:
            text = content.decode('utf-8', errors='ignore')
            output = text_to_word(text)
        return send_file(output, download_name=output_name, as_attachment=True,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        logger.error(f"API xatolik: {e}")
        return jsonify({'error': str(e)[:200]}), 500

@flask_app.route('/api/send-to-bot', methods=['POST'])
def send_to_bot():
    try:
        file = request.files.get('file')
        user_id = request.form.get('user_id')
        if not file or not user_id:
            return jsonify({'error': 'Fayl yoki user_id topilmadi'}), 400
        url = f"https://api.telegram.org/bot{TOKEN}/sendDocument"
        files = {'document': (file.filename, file.read(), file.content_type or 'application/octet-stream')}
        data = {'chat_id': int(user_id), 'caption': '✅ Konvertatsiya natijangiz tayyor!'}
        response = requests.post(url, data=data, files=files, timeout=30)
        if response.status_code == 200:
            return jsonify({'success': True})
        else:
            logger.error(f"Botga yuborishda xatolik: {response.text}")
            return jsonify({'error': 'Botga yuborib bolmadi'}), 500
    except Exception as e:
        logger.error(f"send-to-bot xatolik: {e}")
        return jsonify({'error': str(e)[:200]}), 500

# =============== PRO API ROUTES ===============
@flask_app.route('/api/pro/convert', methods=['POST'])
def api_pro_convert():
    try:
        file = request.files.get('file')
        if not file:
            return jsonify({'error': 'Fayl topilmadi'}), 400
        content = file.read()
        if not file.filename.endswith('.docx'):
            return jsonify({'error': 'Faqat .docx fayllar qabul qilinadi'}), 400
        doc = Document(BytesIO(content))
        pdf_doc = fitz.open()
        pdf_doc.new_page()
        page = pdf_doc[0]
        y = 72
        for para in doc.paragraphs:
            if para.text.strip():
                page.insert_text(fitz.Point(72, y), para.text[:500], fontsize=11)
                y += 14
                if y > 700:
                    pdf_doc.new_page()
                    page = pdf_doc[-1]
                    y = 72
        pdf_output = BytesIO()
        pdf_doc.save(pdf_output)
        pdf_output.seek(0)
        pdf_doc.close()
        output_name = file.filename.rsplit('.', 1)[0] + '.pdf'
        return send_file(pdf_output, download_name=output_name, as_attachment=True, mimetype='application/pdf')
    except Exception as e:
        logger.error(f"Pro convert xatolik: {e}")
        return jsonify({'error': str(e)[:200]}), 500

@flask_app.route('/api/pro/batch', methods=['POST'])
def api_pro_batch():
    try:
        files = request.files.getlist('files')
        if not files or len(files) > 10:
            return jsonify({'error': '1 tadan 10 tagacha fayl yuklash kerak'}), 400
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            for file in files:
                content = file.read()
                ext = file.filename.split('.')[-1].lower()
                output_name = file.filename.rsplit('.', 1)[0] + '.docx'
                if ext == 'pdf':
                    output = pdf_to_word(content, file.filename)
                elif ext in ['xlsx', 'xls', 'csv']:
                    output = excel_to_word(content, ext, file.filename)
                elif ext in ['html', 'htm']:
                    output = html_to_word(content, file.filename)
                elif ext == 'epub':
                    output = epub_to_word(content, file.filename)
                else:
                    text = content.decode('utf-8', errors='ignore')
                    output = text_to_word(text)
                zf.writestr(output_name, output.getvalue())
        zip_buffer.seek(0)
        return send_file(zip_buffer, download_name='converted_files.zip', as_attachment=True, mimetype='application/zip')
    except Exception as e:
        logger.error(f"Batch xatolik: {e}")
        return jsonify({'error': str(e)[:200]}), 500

@flask_app.route('/api/pro/password', methods=['POST'])
def api_pro_password():
    try:
        file = request.files.get('file')
        password = request.form.get('password')
        action = request.form.get('action')
        if not file or not password:
            return jsonify({'error': 'Fayl va parol kiritilishi shart'}), 400
        content = file.read()
        reader = PdfReader(BytesIO(content))
        writer = PdfWriter()
        if action == 'unlock' and reader.is_encrypted:
            try:
                reader.decrypt(password)
            except:
                return jsonify({'error': 'Noto\'g\'ri parol'}), 400
        for page in reader.pages:
            writer.add_page(page)
        if action == 'lock':
            writer.encrypt(password)
        output = BytesIO()
        writer.write(output)
        output.seek(0)
        prefix = 'locked_' if action == 'lock' else 'unlocked_'
        return send_file(output, download_name=prefix + file.filename, as_attachment=True, mimetype='application/pdf')
    except Exception as e:
        logger.error(f"Password xatolik: {e}")
        return jsonify({'error': str(e)[:200]}), 500

def run_flask():
    port = int(os.getenv("PORT", "8080"))
    flask_app.run(host='0.0.0.0', port=port)

# =============== SQLITE ===============
def init_db():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("CREATE TABLE IF NOT EXISTS users (user_id INTEGER PRIMARY KEY, username TEXT, first_seen TEXT)")
    cur.execute("CREATE TABLE IF NOT EXISTS conversions (id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, file_type TEXT, file_name TEXT, success INTEGER, created_at TEXT)")
    cur.execute("CREATE TABLE IF NOT EXISTS pro_users (user_id INTEGER PRIMARY KEY, pro_until TEXT)")
    conn.commit()
    conn.close()

def log_user(user_id: int, username: str):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("INSERT OR IGNORE INTO users (user_id, username, first_seen) VALUES (?, ?, ?)",
                (user_id, username, datetime.now().isoformat()))
    conn.commit()
    conn.close()

def log_conversion(user_id: int, file_type: str, file_name: str, success: bool):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("INSERT INTO conversions (user_id, file_type, file_name, success, created_at) VALUES (?, ?, ?, ?, ?)",
                (user_id, file_type, file_name, int(success), datetime.now().isoformat()))
    conn.commit()
    conn.close()

def get_stats() -> dict:
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM users")
    total_users = cur.fetchone()[0]
    cur.execute("SELECT COUNT(*) FROM conversions")
    total_conversions = cur.fetchone()[0]
    cur.execute("SELECT COUNT(*) FROM conversions WHERE success = 1")
    successful = cur.fetchone()[0]
    cur.execute("SELECT file_type, COUNT(*) c FROM conversions GROUP BY file_type ORDER BY c DESC LIMIT 5")
    top_types = cur.fetchall()
    conn.close()
    return {"total_users": total_users, "total_conversions": total_conversions, "successful": successful, "top_types": top_types}

def is_pro_user(user_id: int) -> bool:
    if datetime.now() < PRO_FREE_UNTIL:
        return True
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT pro_until FROM pro_users WHERE user_id = ?", (user_id,))
    row = cur.fetchone()
    conn.close()
    if row:
        pro_until = datetime.fromisoformat(row[0])
        return pro_until > datetime.now()
    return False

def save_pro_status(user_id: int, pro_until: datetime):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("INSERT OR REPLACE INTO pro_users (user_id, pro_until) VALUES (?, ?)",
                (user_id, pro_until.isoformat()))
    conn.commit()
    conn.close()

# =============== WORD HUJJAT ===============
def create_doc(title: str, original_name: str = "") -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    title_heading = doc.add_heading(title, level=0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_heading.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 51, 102)
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Sana: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    if original_name:
        file_para = doc.add_paragraph()
        file_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        file_run = file_para.add_run(f"Asl fayl: {original_name}")
        file_run.font.size = Pt(9)
        file_run.font.color.rgb = RGBColor(128, 128, 128)
    line_para = doc.add_paragraph()
    line_run = line_para.add_run("─" * 70)
    line_run.font.size = Pt(6)
    line_run.font.color.rgb = RGBColor(200, 200, 200)
    doc.add_paragraph()
    return doc

def clean_text(text: str) -> str:
    if not text or len(text.strip()) < 2:
        return ""
    text = text.strip()
    if not any(c.isalnum() for c in text) and len(text) < 15:
        return ""
    if re.match(r'^[\.\-\s=_~|]{5,}$', text):
        return ""
    text = re.sub(r'[ \t]+', ' ', text)
    return text

def looks_like_heading(block, body_font_size: float = 11.0) -> bool:
    lines = block.get("lines", [])
    if len(lines) != 1:
        return False
    spans = lines[0].get("spans", [])
    text = "".join(s["text"] for s in spans).strip()
    if not text or len(text) > 70:
        return False
    letters = [c for c in text if c.isalpha()]
    if not letters:
        return False
    upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
    max_size = max((s.get("size", body_font_size) for s in spans), default=body_font_size)
    is_bold = any(bool(s.get("flags", 0) & 2 ** 4) for s in spans)
    return upper_ratio > 0.8 or max_size >= body_font_size + 2 or (is_bold and len(text.split()) <= 6)

# =============== PDF → WORD ===============
def pdf_to_word(pdf_content: bytes, original_name: str, progress_cb=None) -> BytesIO:
    doc = create_doc("📄 PDF Hujjat", original_name)
    pdf = fitz.open(stream=pdf_content, filetype="pdf")
    total_pages = len(pdf)
    for page_num in range(total_pages):
        page = pdf[page_num]
        if total_pages > 1:
            marker = doc.add_paragraph()
            marker.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            marker_run = marker.add_run(f"— {page_num + 1} —")
            marker_run.font.size = Pt(7)
            marker_run.font.color.rgb = RGBColor(190, 190, 190)
        table_rects = []
        tabs = None
        try:
            drawings_count = len(page.get_drawings())
            if drawings_count < 800:
                tabs = page.find_tables()
                for table_obj in tabs.tables:
                    table_rects.append(fitz.Rect(table_obj.bbox))
        except:
            pass
        def inside_table(bbox):
            r = fitz.Rect(bbox)
            return any(r.intersects(tr) for tr in table_rects)
        items = []
        try:
            blocks = page.get_text("dict")["blocks"]
            for block in blocks:
                if block.get("type") != 0:
                    continue
                bbox = block.get("bbox", (0, 0, 0, 0))
                if inside_table(bbox):
                    continue
                items.append((bbox[1], bbox[0], "text_block", block))
        except:
            pass
        try:
            for img_index, img in enumerate(page.get_images(full=True), 1):
                xref = img[0]
                try:
                    img_rects = page.get_image_rects(xref)
                    bbox = img_rects[0] if img_rects else fitz.Rect(0, 0, 1, 1)
                    items.append((bbox.y0, bbox.x0, "image", (xref, img_index)))
                except:
                    pass
        except:
            pass
        if tabs is not None:
            for table_obj in tabs.tables:
                bbox = table_obj.bbox
                items.append((bbox[1], bbox[0], "table", table_obj))
        text_x_positions = sorted(it[1] for it in items if it[2] == "text_block")
        column_bounds = []
        if text_x_positions:
            clusters = [[text_x_positions[0]]]
            for x in text_x_positions[1:]:
                if x - clusters[-1][-1] > 40:
                    clusters.append([x])
                else:
                    clusters[-1].append(x)
            clusters = [c for c in clusters if len(c) >= 2]
            if len(clusters) >= 2:
                for c in clusters:
                    column_bounds.append(min(c))
                column_bounds.sort()
        def column_index(x_val):
            if not column_bounds:
                return 0
            best = 0
            best_dist = abs(x_val - column_bounds[0])
            for idx, cb in enumerate(column_bounds):
                d = abs(x_val - cb)
                if d < best_dist:
                    best_dist = d
                    best = idx
            return best
        items.sort(key=lambda it: (column_index(it[1]), round(it[0] / 5)))
        for y, x, kind, payload in items:
            if kind == "text_block":
                block = payload
                if looks_like_heading(block):
                    text = "".join(s["text"] for line in block.get("lines", []) for s in line.get("spans", [])).strip()
                    text = re.sub(r'\s+', ' ', text)
                    if text:
                        doc.add_heading(text, level=2)
                    continue
                para = doc.add_paragraph()
                any_text = False
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        span_text = span["text"]
                        if not span_text.strip():
                            continue
                        cleaned_span = re.sub(r'\s+', ' ', span_text)
                        run = para.add_run(cleaned_span + " ")
                        any_text = True
                        font_size = span.get("size", 11)
                        run.font.size = Pt(max(6, min(36, round(font_size))))
                        font_name = span.get("font", "").lower()
                        flags = span.get("flags", 0)
                        run.font.bold = bool(flags & 2 ** 4) or "bold" in font_name
                        run.font.italic = bool(flags & 2 ** 1) or "italic" in font_name or "oblique" in font_name
                if not any_text:
                    p_elem = para._element
                    p_elem.getparent().remove(p_elem)
            elif kind == "image":
                xref, img_index = payload
                try:
                    base_image = pdf.extract_image(xref)
                    image_bytes = base_image["image"]
                    pil_img = Image.open(BytesIO(image_bytes))
                    if pil_img.width < 40 or pil_img.height < 40:
                        continue
                    doc.add_picture(BytesIO(image_bytes), width=Inches(6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    pass
            elif kind == "table":
                table_obj = payload
                try:
                    table_data = table_obj.extract()
                    if not table_data or len(table_data) < 1:
                        continue
                    has_content = any(cell and str(cell).strip() for row in table_data for cell in row if cell)
                    if not has_content:
                        continue
                    rows = len(table_data)
                    cols = max(len(row) for row in table_data if row)
                    if cols == 0:
                        continue
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Light Grid Accent 1'
                    for i, row in enumerate(table_data):
                        for j in range(cols):
                            if j < len(row) and row[j]:
                                cell_text = clean_text(str(row[j]))
                                if cell_text:
                                    table.cell(i, j).text = cell_text
                    doc.add_paragraph()
                except:
                    pass
        if page_num < total_pages - 1:
            doc.add_page_break()
        if progress_cb:
            try:
                progress_cb(page_num + 1, total_pages)
            except:
                pass
    pdf.close()
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# =============== EXCEL → WORD ===============
def excel_to_word(file_content: bytes, file_ext: str, original_name: str) -> BytesIO:
    import pandas as pd
    doc = create_doc("📊 Excel Ma'lumotlari", original_name)
    try:
        if file_ext == 'csv':
            df = pd.read_csv(BytesIO(file_content), encoding='utf-8')
        else:
            df = pd.read_excel(BytesIO(file_content))
        df = df.dropna(how='all', axis=1)
        df = df.dropna(how='all', axis=0)
        if len(df) > 0:
            info_para = doc.add_paragraph()
            info_run = info_para.add_run(f"Qatorlar: {len(df)} | Ustunlar: {len(df.columns)}")
            info_run.font.size = Pt(10)
            info_run.font.color.rgb = RGBColor(100, 100, 100)
            doc.add_paragraph()
            table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
            table.style = 'Light Grid Accent 1'
            for j, col in enumerate(df.columns):
                cell = table.cell(0, j)
                cell.text = str(col)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
            for i, (_, row) in enumerate(df.iterrows()):
                for j, value in enumerate(row):
                    cell = table.cell(i + 1, j)
                    if pd.notna(value):
                        cell.text = str(value)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
        else:
            doc.add_paragraph("Ma'lumot topilmadi.")
    except Exception as e:
        doc.add_paragraph(f"Xatolik: {str(e)[:100]}")
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# =============== MATN → WORD ===============
def text_to_word(text: str) -> BytesIO:
    doc = create_doc("📝 Matn Hujjati")
    for para in text.split('\n'):
        cleaned = clean_text(para)
        if cleaned:
            doc.add_paragraph(cleaned)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# =============== HTML → WORD ===============
def html_to_word(file_content: bytes, original_name: str) -> BytesIO:
    from bs4 import BeautifulSoup
    import base64
    doc = create_doc("🌐 HTML Hujjat", original_name)
    html = file_content.decode('utf-8', errors='ignore')
    soup = BeautifulSoup(html, 'html.parser')
    for tag in soup(["script", "style"]):
        tag.decompose()
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'li', 'div']):
        text = element.get_text(separator=' ', strip=True)
        cleaned = clean_text(text)
        if not cleaned:
            continue
        if element.name in ('h1', 'h2', 'h3'):
            level = int(element.name[1])
            doc.add_heading(cleaned, level=level)
        else:
            doc.add_paragraph(cleaned)
    for img_tag in soup.find_all('img'):
        src = img_tag.get('src', '')
        if src.startswith('data:image'):
            try:
                header, encoded = src.split(',', 1)
                img_bytes = base64.b64decode(encoded)
                doc.add_picture(BytesIO(img_bytes), width=Inches(5))
            except:
                pass
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# =============== EPUB → WORD ===============
def epub_to_word(file_content: bytes, original_name: str) -> BytesIO:
    from ebooklib import epub
    import ebooklib
    from bs4 import BeautifulSoup
    doc = create_doc("📚 Elektron Kitob", original_name)
    temp_epub = tempfile.NamedTemporaryFile(suffix='.epub', delete=False)
    temp_epub.write(file_content)
    temp_epub.close()
    try:
        book = epub.read_epub(temp_epub.name)
        chapter_num = 0
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                for tag in soup(["script", "style"]):
                    tag.decompose()
                text = soup.get_text()
                if text.strip():
                    chapter_num += 1
                    doc.add_heading(f'Bo\'lim {chapter_num}', level=1)
                    for para in text.split('\n'):
                        cleaned = clean_text(para)
                        if cleaned:
                            doc.add_paragraph(cleaned)
                    doc.add_page_break()
    finally:
        os.unlink(temp_epub.name)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# =============== RASM(LAR) → PDF ===============
def images_to_pdf(images_content: list) -> BytesIO:
    if not images_content:
        raise ValueError("Rasmlar topilmadi")
    pil_images = []
    for content in images_content:
        img = Image.open(BytesIO(content))
        if img.mode == 'RGBA':
            background = Image.new('RGB', img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[3])
            img = background
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        pil_images.append(img)
    output = BytesIO()
    if len(pil_images) == 1:
        pil_images[0].save(output, format='PDF')
    else:
        pil_images[0].save(output, format='PDF', save_all=True, append_images=pil_images[1:])
    output.seek(0)
    return output

# =============== RASM OCR ===============
def preprocess_image_for_ocr(image: Image.Image) -> Image.Image:
    if image.mode != 'L':
        image = image.convert('L')
    image = ImageEnhance.Contrast(image).enhance(2.0)
    image = ImageEnhance.Sharpness(image).enhance(1.5)
    if image.width < 1000:
        scale = 1000 / image.width
        new_size = (int(image.width * scale), int(image.height * scale))
        image = image.resize(new_size, Image.LANCZOS)
    return image

def image_ocr(image_content: bytes) -> str:
    try:
        image = Image.open(BytesIO(image_content))
        image = preprocess_image_for_ocr(image)
        best_text = ""
        for psm in ('6', '3', '11'):
            try:
                text = pytesseract.image_to_string(image, lang=OCR_LANGS, config=f'--psm {psm}')
                if len(text.strip()) > len(best_text.strip()):
                    best_text = text
            except:
                pass
        return best_text.strip()
    except:
        return ""

# =============== TUGMALAR ===============
def result_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([[InlineKeyboardButton("✏️ Nomini o'zgartirish", callback_data="rename")]])

def main_menu_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("📖 Yordam", callback_data="help"),
         InlineKeyboardButton("📋 Formatlar", callback_data="formats")],
        [InlineKeyboardButton("ℹ️ Bot haqida", callback_data="about"),
         InlineKeyboardButton("📩 Murojaat", callback_data="contact_admin")],
        [InlineKeyboardButton("📸 Rasm → PDF", callback_data="rasm_pdf_info"),
         InlineKeyboardButton("📊 Statistika", callback_data="public_stats")],
    ])

def back_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([[InlineKeyboardButton("⬅️ Orqaga", callback_data="menu")]])

# =============== MATNLAR ===============
HELP_TEXT = f"""
📖 <b>YORDAM MARKAZI</b>

━━━━━━━━━━━━━━━━━━━━━

📝 <b>1. MATN</b>
Oddiy matn yozing — Word (.docx) faylga o'tkaziladi.

📄 <b>2. PDF</b>
PDF fayl yuboring. Matn, jadvallar va rasmlar saqlanadi.
• Oddiy PDF: matn + jadvallar
• Skanerlangan PDF: OCR
• Ko'p ustunli: to'g'ri tartibda

📊 <b>3. EXCEL</b>
.xlsx, .xls, .csv yuboring.

🌐 <b>4. HTML</b>
.html fayllar, sarlavhalar saqlanadi.

📚 <b>5. EPUB</b>
Elektron kitoblar, bo'limlar alohida.

🖼 <b>6. RASM OCR</b>
Matnli rasm — OCR → Word

📸 <b>7. RASM → PDF</b>
/rasm_pdf → rasmlar → PDF

✏️ <b>8. NOM O'ZGARTIRISH</b>
Natija ostidagi tugma orqali.

━━━━━━━━━━━━━━━━━━━━━

⚠️ Fayl: 50 MB gacha | ⏱ PDF: 4 daqiqa
📞 Admin: {ADMIN_USERNAME}
"""

FORMATS_TEXT = """
📋 <b>QO'LLAB-QUVVATLANADIGAN FORMATLAR</b>

━━━━━━━━━━━━━━━━━━━━━

📄 <b>Hujjatlar:</b>
• .pdf — PDF (matn + rasm + jadval)
• .xlsx, .xls, .csv — Excel
• .html, .htm — Web sahifalar

📚 <b>Kitoblar:</b>
• .epub

📝 <b>Matn:</b>
• Oddiy matn
• .txt, .md, .json, .xml

🖼 <b>Rasmlar:</b>
• .jpg, .png, .bmp (OCR)
• /rasm_pdf (Rasm → PDF)

━━━━━━━━━━━━━━━━━━━━━

✅ <b>Barchasi → .docx</b>
"""

ABOUT_TEXT = f"""
ℹ️ <b>BOT HAQIDA</b>

━━━━━━━━━━━━━━━━━━━━━

🏷 <b>{BOT_NAME}</b>
📌 <b>Versiya 3.0 Pro</b>

💼 Professional fayl konvertatsiya boti

⭐ <b>Afzalliklari:</b>
• PDF ichidagi rasmlar saqlanadi
• Shrift formatlash (bold, italic)
• Ko'p ustunli PDF
• Sarlavhalarni aniqlash
• Real-time progress-bar

━━━━━━━━━━━━━━━━━━━━━

👨‍💼 <b>Admin:</b> {ADMIN_USERNAME}
📞 <b>Aloqa:</b> {ADMIN_USERNAME}
"""

# =============== YUBORISH FUNKSIYASI ===============
async def send_result_document(update: Update, context: ContextTypes.DEFAULT_TYPE, output: BytesIO, filename: str, caption: str):
    data = output.getvalue()
    context.user_data['last_output_bytes'] = data
    context.user_data['last_output_name'] = filename
    await update.message.reply_document(
        document=BytesIO(data),
        filename=filename,
        caption=caption,
        parse_mode='HTML',
        reply_markup=result_keyboard()
    )

def make_progress_bar(current: int, total: int, width: int = 15) -> str:
    if total <= 0:
        return ""
    ratio = current / total
    filled = int(ratio * width)
    bar = "🟦" * filled + "⬜" * (width - filled)
    percent = int(ratio * 100)
    return f"{bar}\n📄 <b>{current}/{total}</b> sahifa (<b>{percent}%</b>)"

# =============== HANDLERLAR ===============

async def start_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    log_user(user.id, user.username or user.first_name)
    message = f"""
✨ <b>Assalomu alaykum, {user.first_name}!</b>

<b>🏷 {BOT_NAME}</b>

━━━━━━━━━━━━━━━━━━━━━

📌 <b>Word (.docx) ga o'tkazadi:</b>

📄 PDF <i>(matn + rasm + jadvallar)</i>
📊 Excel jadvallari
🌐 HTML sahifalar
📚 EPUB kitoblar
📝 Matn fayllari
🖼 Rasmlar <i>(OCR)</i>
📸 Rasmlarni PDF qilish

━━━━━━━━━━━━━━━━━━━━━

⚡ <b>Foydalanish:</b>
1️⃣ Fayl yoki matn yuboring
2️⃣ Avtomatik konvertatsiya
3️⃣ Tayyor faylni yuklab oling

📌 Maksimal hajm: <b>50 MB</b>
📌 <b>24/7</b> ishlaydi
📌 <b>Bepul</b> ✅

💎 <b>/pro</b> — Pro funksiyalar

👨‍💼 Admin: {ADMIN_USERNAME}

Quyidagi menyudan foydalaning 👇
"""
    await update.message.reply_text(message, parse_mode='HTML', reply_markup=main_menu_keyboard())


async def menu_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    if query.data == "help":
        await query.edit_message_text(HELP_TEXT, parse_mode='HTML', reply_markup=back_keyboard())
    elif query.data == "formats":
        await query.edit_message_text(FORMATS_TEXT, parse_mode='HTML', reply_markup=back_keyboard())
    elif query.data == "about":
        await query.edit_message_text(ABOUT_TEXT, parse_mode='HTML', reply_markup=back_keyboard())
    elif query.data == "menu":
        user = query.from_user
        message = f"""✨ <b>{user.first_name}</b>, asosiy menyu\n\n💼 <b>{BOT_NAME}</b>\n\nFayl yoki matn yuboring — avtomatik Word'ga o'tkazaman.\nYoki quyidagi tugmalardan foydalaning 👇"""
        await query.edit_message_text(message, parse_mode='HTML', reply_markup=main_menu_keyboard())
    elif query.data == "rasm_pdf_info":
        message = """📸 <b>RASM → PDF</b>\n\n━━━━━━━━━━━━━━━━━━━━━\n\nBir nechta rasmni bitta PDF qilish:\n\n1️⃣ /rasm_pdf buyrug'ini yuboring\n2️⃣ Rasmlarni birin-ketin jo'nating\n3️⃣ "✅ PDF QILISH" tugmasini bosing\n\n📌 Rasmlar A4 formatda, markazda joylashadi."""
        await query.edit_message_text(message, parse_mode='HTML', reply_markup=back_keyboard())
    elif query.data == "public_stats":
        stats = get_stats()
        message = f"""📊 <b>BOT STATISTIKASI</b>\n\n━━━━━━━━━━━━━━━━━━━━━\n\n👥 Foydalanuvchilar: <b>{stats['total_users']}</b>\n🔄 Konvertatsiyalar: <b>{stats['total_conversions']}</b>\n✅ Muvaffaqiyatli: <b>{stats['successful']}</b>\n\n━━━━━━━━━━━━━━━━━━━━━\n\n🌟 24/7 ishlaydi\n📌 Bepul xizmat\n⚡ Tez va sifatli"""
        await query.edit_message_text(message, parse_mode='HTML', reply_markup=back_keyboard())
    elif query.data == "contact_admin":
        context.user_data['awaiting_contact_message'] = True
        await query.message.reply_text("📩 <b>Murojaat yuborish</b>\n\nXabaringizni yozing — to'g'ridan-to'g'ri adminga yuboriladi.\n❌ Bekor qilish uchun /cancel", parse_mode='HTML')
    elif query.data == "rename":
        if not context.user_data.get('last_output_bytes'):
            await query.message.reply_text("⚠️ O'zgartiriladigan fayl topilmadi.")
            return
        context.user_data['awaiting_rename'] = True
        await query.message.reply_text("✏️ <b>Yangi fayl nomini yozing</b> (kengaytmasiz):", parse_mode='HTML')
    elif query.data == "make_pdf":
        images = context.user_data.get('pdf_images', [])
        if not images:
            await query.message.reply_text("⚠️ Hech qanday rasm topilmadi. Avval rasm yuboring.")
            return
        await query.edit_message_reply_markup(reply_markup=None)
        progress = await query.message.reply_text(f"⚙️ <b>{len(images)}</b> ta rasmdan PDF yaratilmoqda...", parse_mode='HTML')
        try:
            output = images_to_pdf(images)
            await progress.delete()
            data = output.getvalue()
            context.user_data['last_output_bytes'] = data
            context.user_data['last_output_name'] = "rasmlar.pdf"
            await query.message.reply_document(document=BytesIO(data), filename="rasmlar.pdf", caption=f"✅ <b>Tayyor!</b>\n🖼 <b>{len(images)}</b> ta rasmdan PDF yaratildi.", parse_mode='HTML', reply_markup=result_keyboard())
            log_conversion(query.from_user.id, "images_to_pdf", "rasmlar.pdf", success=True)
        except Exception as e:
            await progress.delete()
            logger.error(f"Rasm→PDF xatolik: {e}")
            await query.message.reply_text(f"❌ Xatolik: {str(e)[:150]}\n📞 {ADMIN_USERNAME}")
        finally:
            context.user_data['collecting_pdf'] = False
            context.user_data['pdf_images'] = []
    elif query.data == "cancel_pdf":
        context.user_data['collecting_pdf'] = False
        context.user_data['pdf_images'] = []
        await query.edit_message_reply_markup(reply_markup=None)
        await query.message.reply_text("❌ Bekor qilindi.")


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(HELP_TEXT, parse_mode='HTML', reply_markup=back_keyboard())

async def formats_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(FORMATS_TEXT, parse_mode='HTML', reply_markup=back_keyboard())

async def admin_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID:
        await update.message.reply_text("⛔ <b>Ruxsat yo'q!</b>\nFaqat admin uchun.", parse_mode='HTML')
        return
    stats = get_stats()
    message = f"""👑 <b>ADMIN PANEL</b>\n\n━━━━━━━━━━━━━━━━━━━━━\n\n🟢 <b>Holat:</b> Aktiv\n📌 <b>Versiya:</b> 3.0 Pro\n\n━━━━━━━━━━━━━━━━━━━━━\n\n📊 <b>STATISTIKA:</b>\n\n👥 Foydalanuvchilar: <b>{stats['total_users']}</b>\n🔄 Jami: <b>{stats['total_conversions']}</b>\n✅ Muvaffaqiyatli: <b>{stats['successful']}</b>\n❌ Xatolik: <b>{stats['total_conversions'] - stats['successful']}</b>\n\n━━━━━━━━━━━━━━━━━━━━━\n\n📈 <b>ENG KO'P:</b>\n"""
    emoji_map = {"pdf": "📄", "xlsx": "📊", "csv": "📊", "xls": "📊", "text": "📝", "photo_ocr": "🖼", "html": "🌐", "epub": "📚", "images_to_pdf": "📸"}
    for ftype, count in stats.get("top_types", [])[:5]:
        emoji = emoji_map.get(ftype, "📎")
        message += f"{emoji} {ftype}: <b>{count}</b>\n"
    message += f"\n👨‍💼 Admin: {ADMIN_USERNAME}"
    await update.message.reply_text(message, parse_mode='HTML')

async def stats_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID:
        await update.message.reply_text("⛔ <b>Ruxsat yo'q!</b>", parse_mode='HTML')
        return
    await admin_command(update, context)

async def about_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(ABOUT_TEXT, parse_mode='HTML', reply_markup=back_keyboard())


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    document = update.message.document
    file_name = document.file_name
    user = update.effective_user
    log_user(user.id, user.username or user.first_name)
    if document.file_size > MAX_FILE_SIZE:
        await update.message.reply_text("❌ Fayl hajmi 50 MB dan oshmasligi kerak.")
        return
    file_ext = file_name.lower().split('.')[-1] if '.' in file_name else ''
    supported = ['pdf', 'xlsx', 'xls', 'csv', 'html', 'htm', 'epub', 'txt', 'md', 'json', 'xml']
    if file_ext not in supported:
        await update.message.reply_text(f"❌ .{file_ext} qo'llab-quvvatlanmaydi.\n📋 /formats")
        return
    file_emoji = {"pdf": "📄", "xlsx": "📊", "xls": "📊", "csv": "📊", "html": "🌐", "htm": "🌐", "epub": "📚"}.get(file_ext, "📎")
    msg = await update.message.reply_text(f"{file_emoji} <b>{file_name}</b> yuklab olinmoqda...\n📦 {document.file_size // 1024} KB", parse_mode='HTML')
    try:
        file = await document.get_file()
        content = await file.download_as_bytearray()
        output_name = file_name.rsplit('.', 1)[0] + '.docx'
        if file_ext == 'pdf':
            loop = asyncio.get_running_loop()
            last_edit = {"time": datetime.min, "text": ""}
            async def update_progress_msg(text: str):
                now = datetime.now()
                if text == last_edit["text"] or (now - last_edit["time"]).total_seconds() < 2:
                    return
                last_edit["time"] = now
                last_edit["text"] = text
                try:
                    await msg.edit_text(text, parse_mode='HTML')
                except:
                    pass
            def sync_progress(current: int, total: int):
                bar = make_progress_bar(current, total)
                text = f"⚙️ <b>{file_name}</b> konvertatsiya qilinmoqda...\n{bar}"
                asyncio.run_coroutine_threadsafe(update_progress_msg(text), loop)
            await msg.edit_text(f"⚙️ <b>{file_name}</b> konvertatsiya qilinmoqda...\n{make_progress_bar(0, 1)}", parse_mode='HTML')
            try:
                output = await asyncio.wait_for(asyncio.to_thread(pdf_to_word, bytes(content), file_name, sync_progress), timeout=PDF_TIMEOUT_SECONDS)
            except asyncio.TimeoutError:
                await msg.delete()
                await update.message.reply_text(f"⏱ <b>{file_name}</b> konvertatsiyasi {PDF_TIMEOUT_SECONDS} soniyadan ko'p vaqt oldi.\n💡 Faylni kichikroq qismlarga bo'lib yuboring.", parse_mode='HTML')
                return
        elif file_ext in ['xlsx', 'xls', 'csv']:
            await msg.edit_text(f"⚙️ <b>{file_name}</b> konvertatsiya qilinmoqda...", parse_mode='HTML')
            output = await asyncio.to_thread(excel_to_word, bytes(content), file_ext, file_name)
        elif file_ext in ['html', 'htm']:
            await msg.edit_text(f"⚙️ <b>{file_name}</b> konvertatsiya qilinmoqda...", parse_mode='HTML')
            output = await asyncio.to_thread(html_to_word, bytes(content), file_name)
        elif file_ext == 'epub':
            await msg.edit_text(f"⚙️ <b>{file_name}</b> konvertatsiya qilinmoqda...", parse_mode='HTML')
            output = await asyncio.to_thread(epub_to_word, bytes(content), file_name)
        else:
            await msg.edit_text(f"⚙️ <b>{file_name}</b> konvertatsiya qilinmoqda...", parse_mode='HTML')
            text = content.decode('utf-8', errors='ignore')
            output = await asyncio.to_thread(text_to_word, text)
        await msg.delete()
        await send_result_document(update, context, output, output_name, f"✅ <b>Tayyor!</b>\n\n{file_emoji} <b>{file_name}</b>\n⬇️ <b>{output_name}</b>\n📅 {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        log_conversion(user.id, file_ext, file_name, success=True)
    except Exception as e:
        await msg.delete()
        logger.error(f"Xatolik ({file_name}): {e}")
        await update.message.reply_text(f"❌ Xatolik yuz berdi.\n📞 Admin: {ADMIN_USERNAME}")
        log_conversion(user.id, file_ext, file_name, success=False)


async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    log_user(user.id, user.username or user.first_name)
    photo = update.message.photo[-1]
    file = await photo.get_file()
    content = await file.download_as_bytearray()
    if context.user_data.get('collecting_pdf'):
        context.user_data.setdefault('pdf_images', []).append(bytes(content))
        count = len(context.user_data['pdf_images'])
        message = f"📥 <b>{count}-rasm qabul qilindi!</b>\n\n📸 Jami: <b>{count}</b> ta rasm\nYana rasm yuboring yoki tugmani bosing 👇"
        keyboard = InlineKeyboardMarkup([[InlineKeyboardButton(f"✅ PDF QILISH ({count} ta)", callback_data="make_pdf")], [InlineKeyboardButton("❌ Bekor qilish", callback_data="cancel_pdf")]])
        await update.message.reply_text(message, parse_mode='HTML', reply_markup=keyboard)
        return
    msg = await update.message.reply_text("🖼 <b>Rasm qayta ishlanmoqda</b> (OCR)...", parse_mode='HTML')
    try:
        text = await asyncio.to_thread(image_ocr, bytes(content))
        if text and len(text.strip()) > 10:
            output = await asyncio.to_thread(text_to_word, text)
            await msg.delete()
            await send_result_document(update, context, output, "rasmdagi_matn.docx", f"✅ <b>Tayyor!</b>\n📏 <b>{len(text)}</b> belgi")
            log_conversion(user.id, "photo_ocr", "photo.jpg", success=True)
        else:
            await msg.delete()
            await update.message.reply_text("⚠️ <b>Rasmdan matn topilmadi.</b>\n\n📌 Sabablar:\n• Rasmda matn yo'q\n• Rasm sifatsiz\n• Qo'l yozuvi\n\n💡 Aniq, bosma matnli rasm yuboring.\n📄 Yoki /rasm_pdf orqali PDF qiling.", parse_mode='HTML')
            log_conversion(user.id, "photo_ocr", "photo.jpg", success=False)
    except Exception as e:
        await msg.delete()
        logger.error(f"Rasm xatolik: {e}")
        await update.message.reply_text(f"❌ Xatolik.\n📞 {ADMIN_USERNAME}")


async def rasm_pdf_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['collecting_pdf'] = True
    context.user_data['pdf_images'] = []
    message = """📸 <b>RASM → PDF REJIMI</b>\n\n━━━━━━━━━━━━━━━━━━━━━\n\n📤 Endi rasmlarni birin-ketin yuboring.\nBarcha rasmlar <b>bitta PDF</b> faylga birlashtiriladi.\n\n📌 <b>Eslatma:</b>\n• Rasmlar yuborgan tartibda joylashadi\n• Har bir rasm alohida sahifada\n• A4 formatda, markazda\n\n❌ Bekor qilish uchun /cancel"""
    await update.message.reply_text(message, parse_mode='HTML')


async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text
    user = update.effective_user
    log_user(user.id, user.username or user.first_name)
    if context.user_data.get('awaiting_contact_message'):
        context.user_data['awaiting_contact_message'] = False
        username_display = f"@{user.username}" if user.username else user.first_name
        try:
            await context.bot.send_message(ADMIN_ID, f"📩 <b>YANGI MUROJAAT</b>\n\n👤 {username_display}\n🆔 ID: {user.id}\n\n💬 {text}\n\n↩️ Javob: /reply {user.id} <matn>", parse_mode='HTML')
            await update.message.reply_text("✅ <b>Murojaatingiz adminga yuborildi.</b>", parse_mode='HTML', reply_markup=back_keyboard())
        except:
            await update.message.reply_text(f"❌ Yuborib bo'lmadi.\n📞 Admin: {ADMIN_USERNAME}")
        return
    if context.user_data.get('awaiting_rename'):
        context.user_data['awaiting_rename'] = False
        data = context.user_data.get('last_output_bytes')
        old_name = context.user_data.get('last_output_name', 'fayl.docx')
        if not data:
            await update.message.reply_text("⚠️ Fayl topilmadi.")
            return
        ext = old_name.rsplit('.', 1)[-1] if '.' in old_name else 'docx'
        new_name = re.sub(r'[\\/:*?"<>|]', '_', text.strip())
        if not new_name.lower().endswith('.' + ext):
            new_name = f"{new_name}.{ext}"
        context.user_data['last_output_name'] = new_name
        await update.message.reply_document(document=BytesIO(data), filename=new_name, caption=f"✅ <b>Nomi o'zgartirildi:</b>\n📄 {new_name}", parse_mode='HTML', reply_markup=result_keyboard())
        return
    if len(text) < 10:
        await update.message.reply_text("⚠️ Kamida 10 ta belgi kerak.")
        return
    msg = await update.message.reply_text("⏳ <b>Matn qayta ishlanmoqda...</b>", parse_mode='HTML')
    try:
        output = await asyncio.to_thread(text_to_word, text)
        await msg.delete()
        await send_result_document(update, context, output, "matn.docx", f"✅ <b>Tayyor!</b>\n📏 <b>{len(text)}</b> belgi")
        log_conversion(user.id, "text", "matn.txt", success=True)
    except Exception as e:
        await msg.delete()
        logger.error(f"Matn xatolik: {e}")
        await update.message.reply_text("❌ Xatolik yuz berdi.")


async def cancel_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['awaiting_contact_message'] = False
    context.user_data['awaiting_rename'] = False
    context.user_data['collecting_pdf'] = False
    await update.message.reply_text("❌ <b>Bekor qilindi.</b>", parse_mode='HTML', reply_markup=main_menu_keyboard())


async def reply_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID:
        await update.message.reply_text("⛔ Faqat admin uchun.")
        return
    args = context.args
    if len(args) < 2:
        await update.message.reply_text("/reply <user_id> <matn>")
        return
    try:
        target_id = int(args[0])
    except ValueError:
        await update.message.reply_text("⚠️ user_id raqam bo'lishi kerak.")
        return
    reply_text = " ".join(args[1:])
    try:
        await context.bot.send_message(target_id, f"📩 <b>ADMIN JAVOBI:</b>\n\n{reply_text}", parse_mode='HTML')
        await update.message.reply_text("✅ Javob yuborildi.")
    except Exception as e:
        await update.message.reply_text(f"❌ Yuborib bo'lmadi: {str(e)[:150]}")


async def pro_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    log_user(user.id, user.username or user.first_name)
    if is_pro_user(user.id):
        pro_url = os.getenv("WEBHOOK_URL", "https://worker-production-019c.up.railway.app") + "/pro"
        if datetime.now() < PRO_FREE_UNTIL:
            days = (PRO_FREE_UNTIL - datetime.now()).days
            status = f"🆓 Bepul sinov: {days} kun qoldi"
        else:
            status = "💎 Pro faol"
        keyboard = InlineKeyboardMarkup([
            [InlineKeyboardButton("🚀 Pro Web App", web_app={"url": pro_url})],
            [InlineKeyboardButton("🔙 Orqaga", callback_data="menu")],
        ])
        await update.message.reply_text(f"💎 <b>PRO VERSIYA</b>\n\n{status}\n\n🔥 Word→PDF | Batch+ZIP | Parol", parse_mode='HTML', reply_markup=keyboard)
    else:
        await update.message.reply_text(f"🔒 Pro yopiq.\n📞 {ADMIN_USERNAME}")


async def givepro_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != ADMIN_ID:
        await update.message.reply_text("⛔ Faqat admin uchun.")
        return
    args = context.args
    if len(args) < 2:
        await update.message.reply_text("/givepro <user_id> <kun>")
        return
    try:
        user_id = int(args[0])
        days = int(args[1])
    except:
        await update.message.reply_text("Raqam kiriting.")
        return
    save_pro_status(user_id, datetime.now() + timedelta(days=days))
    await update.message.reply_text(f"✅ {user_id} ga {days} kunlik Pro berildi!")
    try:
        await context.bot.send_message(user_id, f"🎉 Sizga {days} kunlik PRO status berildi!\n/pro orqali kiring!")
    except:
        pass


async def error_handler(update: object, context: ContextTypes.DEFAULT_TYPE):
    logger.error(f"Xatolik: {context.error}")


# =============== MAIN ===============
def main():
    logger.info(f"🤖 {BOT_NAME} ishga tushmoqda...")
    init_db()
    flask_thread = threading.Thread(target=run_flask, daemon=True)
    flask_thread.start()
    logger.info("🌐 Flask Web App server ishga tushdi")
    app = Application.builder().token(TOKEN).build()
    app.add_handler(CommandHandler("start", start_command))
    app.add_handler(CommandHandler("help", help_command))
    app.add_handler(CommandHandler("formats", formats_command))
    app.add_handler(CommandHandler("admin", admin_command))
    app.add_handler(CommandHandler("stats", stats_command))
    app.add_handler(CommandHandler("about", about_command))
    app.add_handler(CommandHandler("rasm_pdf", rasm_pdf_command))
    app.add_handler(CommandHandler("cancel", cancel_command))
    app.add_handler(CommandHandler("reply", reply_command))
    app.add_handler(CommandHandler("pro", pro_command))
    app.add_handler(CommandHandler("givepro", givepro_command))
    app.add_handler(CallbackQueryHandler(menu_callback))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    app.add_error_handler(error_handler)
    WEBHOOK_URL = os.getenv("WEBHOOK_URL", "")
    PORT = int(os.getenv("PORT", "8080"))
    if WEBHOOK_URL:
        logger.info(f"🌐 Webhook: {WEBHOOK_URL}")
        app.run_webhook(listen="0.0.0.0", port=PORT, webhook_url=f"{WEBHOOK_URL}/webhook", drop_pending_updates=True)
    else:
        logger.info("📡 Polling rejimi")
        app.run_polling(drop_pending_updates=True)

if __name__ == "__main__":
    main()
